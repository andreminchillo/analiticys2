"""
Gerador de Relatórios em Word
"""
import os
from typing import Dict, List
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn


class WordReportGenerator:
    def __init__(self):
        pass
    
    def generate_comprehensive_report(self, vendor_data: Dict, output_path: str) -> str:
        """
        Gera relatório completo em Word com dados de todos os vendedores
        
        Args:
            vendor_data: Dados processados de todos os vendedores
            output_path: Caminho para salvar o arquivo
            
        Returns:
            Caminho do arquivo gerado
        """
        try:
            print("📄 Gerando relatório completo em Word...")
            
            # Cria documento
            doc = Document()
            
            # Configura margens
            sections = doc.sections
            for section in sections:
                section.top_margin = Inches(1)
                section.bottom_margin = Inches(1)
                section.left_margin = Inches(1)
                section.right_margin = Inches(1)
            
            # Adiciona conteúdo
            self._add_title_page(doc, vendor_data)
            self._add_executive_summary(doc, vendor_data)
            self._add_team_statistics(doc, vendor_data)
            self._add_vendor_reports(doc, vendor_data)
            self._add_recommendations(doc, vendor_data)
            
            # Salva documento
            doc.save(output_path)
            print(f"✅ Relatório salvo em: {output_path}")
            
            return output_path
            
        except Exception as e:
            print(f"❌ Erro ao gerar relatório: {str(e)}")
            raise
    
    def _add_title_page(self, doc: Document, vendor_data: Dict):
        """
        Adiciona página de título
        """
        # Título principal
        title = doc.add_heading('RELATÓRIO DE ANÁLISE DE VENDAS VVN', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Subtítulo
        subtitle = doc.add_heading('Análise Completa de Performance por Vendedor', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Informações do relatório
        doc.add_paragraph()  # Espaço
        
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = info_para.add_run(f"Data do Relatório: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        run.bold = True
        
        info_para.add_run(f"\\nTotal de Vendedores: {vendor_data.get('total_vendedores', 0)}")
        info_para.add_run(f"\\nTotal de Ligações Analisadas: {vendor_data.get('total_conversas', 0)}")
        
        # Quebra de página
        doc.add_page_break()
    
    def _add_executive_summary(self, doc: Document, vendor_data: Dict):
        """
        Adiciona resumo executivo
        """
        doc.add_heading('📊 RESUMO EXECUTIVO', level=1)
        
        team_stats = vendor_data.get('estatisticas_equipe', {})
        
        # Parágrafo introdutório
        intro_para = doc.add_paragraph()
        intro_para.add_run("Este relatório apresenta uma análise abrangente da performance de vendas da equipe VVN, ").add_run(
            f"baseada na análise de {vendor_data.get('total_conversas', 0)} ligações "
        ).bold = True
        intro_para.add_run(f"de {vendor_data.get('total_vendedores', 0)} vendedores diferentes.")
        
        # Métricas principais
        doc.add_heading('🎯 Métricas Principais da Equipe', level=2)
        
        metrics_table = doc.add_table(rows=4, cols=2)
        metrics_table.style = 'Table Grid'
        
        metrics_data = [
            ('📈 Nota Média da Equipe', f"{team_stats.get('nota_media_equipe', 0)}/10"),
            ('💰 Taxa de Conversão Média', f"{team_stats.get('taxa_conversao_media', 0)}%"),
            ('😊 Sentimento Positivo Médio', f"{team_stats.get('sentimento_positivo_medio', 0)}%"),
            ('🏆 Melhor Vendedor', team_stats.get('melhor_vendedor', 'N/A'))
        ]
        
        for i, (metric, value) in enumerate(metrics_data):
            metrics_table.cell(i, 0).text = metric
            metrics_table.cell(i, 1).text = str(value)
            metrics_table.cell(i, 0).paragraphs[0].runs[0].bold = True
        
        # Resultados consolidados
        doc.add_heading('📋 Resultados Consolidados', level=2)
        
        results_para = doc.add_paragraph()
        results_para.add_run("Vendas Fechadas: ").bold = True
        results_para.add_run(f"{team_stats.get('total_vendas_fechadas', 0)} ligações")
        
        results_para.add_run("\\nFollow-ups Agendados: ").bold = True
        results_para.add_run(f"{team_stats.get('total_follow_ups', 0)} ligações")
        
        results_para.add_run("\\nClientes Perdidos: ").bold = True
        results_para.add_run(f"{team_stats.get('total_clientes_perdidos', 0)} ligações")
        
        # Vendedor que precisa de atenção
        if team_stats.get('vendedor_precisa_atencao'):
            attention_para = doc.add_paragraph()
            attention_para.add_run("⚠️ Vendedor que Precisa de Atenção: ").bold = True
            attention_para.add_run(team_stats['vendedor_precisa_atencao'])
        
        doc.add_page_break()
    
    def _add_team_statistics(self, doc: Document, vendor_data: Dict):
        """
        Adiciona estatísticas da equipe
        """
        doc.add_heading('📈 ESTATÍSTICAS DA EQUIPE', level=1)
        
        # Ranking de vendedores
        doc.add_heading('🏆 Ranking de Vendedores', level=2)
        
        vendors = vendor_data.get('vendedores', {})
        
        # Ordena vendedores por nota média
        sorted_vendors = sorted(
            vendors.items(),
            key=lambda x: x[1]['estatisticas'].get('nota_media', 0),
            reverse=True
        )
        
        if sorted_vendors:
            ranking_table = doc.add_table(rows=len(sorted_vendors) + 1, cols=5)
            ranking_table.style = 'Table Grid'
            
            # Cabeçalho
            headers = ['Posição', 'Vendedor', 'Nota Média', 'Taxa Conversão', 'Total Ligações']
            for i, header in enumerate(headers):
                cell = ranking_table.cell(0, i)
                cell.text = header
                cell.paragraphs[0].runs[0].bold = True
            
            # Dados dos vendedores
            for i, (vendor_name, vendor_info) in enumerate(sorted_vendors, 1):
                stats = vendor_info['estatisticas']
                
                ranking_table.cell(i, 0).text = str(i)
                ranking_table.cell(i, 1).text = vendor_name
                ranking_table.cell(i, 2).text = f"{stats.get('nota_media', 0)}/10"
                ranking_table.cell(i, 3).text = f"{stats.get('taxa_conversao', 0)}%"
                ranking_table.cell(i, 4).text = str(stats.get('total_conversas', 0))
        
        doc.add_page_break()
    
    def _add_vendor_reports(self, doc: Document, vendor_data: Dict):
        """
        Adiciona relatórios individuais por vendedor
        """
        doc.add_heading('👥 RELATÓRIOS INDIVIDUAIS POR VENDEDOR', level=1)
        
        vendors = vendor_data.get('vendedores', {})
        
        for vendor_name, vendor_info in vendors.items():
            self._add_individual_vendor_report(doc, vendor_name, vendor_info)
            doc.add_page_break()
    
    def _add_individual_vendor_report(self, doc: Document, vendor_name: str, vendor_info: Dict):
        """
        Adiciona relatório individual de um vendedor
        """
        # Título do vendedor
        vendor_title = doc.add_heading(f'👤 {vendor_name}', level=2)
        
        stats = vendor_info.get('estatisticas', {})
        insights = vendor_info.get('insights', {})
        conversations = vendor_info.get('conversas', [])
        
        # Resumo do vendedor
        doc.add_heading('📊 Resumo de Performance', level=3)
        
        summary_table = doc.add_table(rows=6, cols=2)
        summary_table.style = 'Table Grid'
        
        summary_data = [
            ('🎯 Nível de Performance', insights.get('nivel_performance', 'N/A')),
            ('📈 Nota Geral', f"{insights.get('nota_geral', 0)}/10"),
            ('📞 Total de Ligações', str(insights.get('total_ligacoes', 0))),
            ('💰 Taxa de Conversão', f"{insights.get('taxa_conversao', 0)}%"),
            ('😊 Taxa Sentimento Positivo', f"{insights.get('taxa_sentimento_positivo', 0)}%"),
            ('⭐ Taxa Ligações Qualidade', f"{insights.get('taxa_ligacoes_qualidade', 0)}%")
        ]
        
        for i, (metric, value) in enumerate(summary_data):
            summary_table.cell(i, 0).text = metric
            summary_table.cell(i, 1).text = str(value)
            summary_table.cell(i, 0).paragraphs[0].runs[0].bold = True
        
        # Pontos fortes
        pontos_fortes = insights.get('pontos_fortes_principais', [])
        if pontos_fortes:
            doc.add_heading('✅ Pontos Fortes Principais', level=3)
            for ponto in pontos_fortes:
                doc.add_paragraph(f"• {ponto}", style='List Bullet')
        
        # Áreas de melhoria
        areas_melhoria = insights.get('areas_melhoria_principais', [])
        if areas_melhoria:
            doc.add_heading('🔧 Áreas de Melhoria', level=3)
            for area in areas_melhoria:
                doc.add_paragraph(f"• {area}", style='List Bullet')
        
        # Prioridades de treinamento
        prioridades = insights.get('prioridades_treinamento', [])
        if prioridades:
            doc.add_heading('🎓 Prioridades de Treinamento', level=3)
            for prioridade in prioridades:
                doc.add_paragraph(f"• {prioridade}", style='List Bullet')
        
        # Objeções mais enfrentadas
        objecoes = insights.get('objecoes_mais_enfrentadas', [])
        if objecoes:
            doc.add_heading('❓ Objeções Mais Enfrentadas', level=3)
            for objecao in objecoes:
                doc.add_paragraph(f"• {objecao}", style='List Bullet')
        
        # Produtos mais trabalhados
        produtos = insights.get('produtos_mais_trabalhados', [])
        if produtos:
            doc.add_heading('🛍️ Produtos Mais Trabalhados', level=3)
            for produto in produtos:
                doc.add_paragraph(f"• {produto}", style='List Bullet')
        
        # Recomendação principal
        recomendacao = insights.get('recomendacao_principal')
        if recomendacao:
            doc.add_heading('💡 Recomendação Principal', level=3)
            doc.add_paragraph(recomendacao)
        
        # Próximos passos
        proximos_passos = insights.get('proximos_passos', [])
        if proximos_passos:
            doc.add_heading('🚀 Próximos Passos', level=3)
            for passo in proximos_passos:
                doc.add_paragraph(f"• {passo}", style='List Bullet')
        
        # Estatísticas detalhadas
        doc.add_heading('📋 Estatísticas Detalhadas', level=3)
        
        detailed_stats = doc.add_paragraph()
        detailed_stats.add_run("Nota Média: ").bold = True
        detailed_stats.add_run(f"{stats.get('nota_media', 0)}/10")
        
        detailed_stats.add_run(" | Melhor Nota: ").bold = True
        detailed_stats.add_run(f"{stats.get('melhor_nota', 0)}/10")
        
        detailed_stats.add_run(" | Pior Nota: ").bold = True
        detailed_stats.add_run(f"{stats.get('pior_nota', 0)}/10")
        
        detailed_stats.add_run("\\nVendas Fechadas: ").bold = True
        detailed_stats.add_run(f"{stats.get('vendas_fechadas', 0)}")
        
        detailed_stats.add_run(" | Follow-ups: ").bold = True
        detailed_stats.add_run(f"{stats.get('follow_ups_agendados', 0)}")
        
        detailed_stats.add_run(" | Clientes Perdidos: ").bold = True
        detailed_stats.add_run(f"{stats.get('clientes_perdidos', 0)}")
        
        # Distribuição de classificações
        classificacoes = stats.get('classificacoes', {})
        if classificacoes:
            doc.add_heading('🏷️ Distribuição de Classificações', level=3)
            
            class_para = doc.add_paragraph()
            class_para.add_run(f"A (Excelente): {classificacoes.get('A', 0)} ligações | ")
            class_para.add_run(f"B (Boa): {classificacoes.get('B', 0)} ligações | ")
            class_para.add_run(f"C (Regular): {classificacoes.get('C', 0)} ligações | ")
            class_para.add_run(f"D (Precisa Melhorar): {classificacoes.get('D', 0)} ligações")
        
        # Lista de ligações (resumida)
        if conversations:
            doc.add_heading('📞 Resumo das Ligações', level=3)
            
            for i, conversation in enumerate(conversations[:5], 1):  # Máximo 5 ligações
                analysis = conversation.get('analise', {})
                classification = conversation.get('classificacao', {})
                
                call_para = doc.add_paragraph()
                call_para.add_run(f"Ligação {i}: ").bold = True
                call_para.add_run(f"{analysis.get('arquivo_origem', 'N/A')}")
                
                call_para.add_run(f" | Nota: {classification.get('nota_vendedor', 0)}/10")
                call_para.add_run(f" | Classificação: {classification.get('classificacao', 'N/A')}")
                call_para.add_run(f" | Resultado: {analysis.get('resultado_conversa', 'N/A')}")
                
                if analysis.get('resumo_executivo'):
                    call_para.add_run(f"\\n   Resumo: {analysis['resumo_executivo']}")
            
            if len(conversations) > 5:
                doc.add_paragraph(f"... e mais {len(conversations) - 5} ligação(ões)")
    
    def _add_recommendations(self, doc: Document, vendor_data: Dict):
        """
        Adiciona recomendações gerais
        """
        doc.add_heading('💡 RECOMENDAÇÕES GERAIS', level=1)
        
        # Recomendações para a equipe
        doc.add_heading('👥 Para a Equipe', level=2)
        
        team_stats = vendor_data.get('estatisticas_equipe', {})
        vendors = vendor_data.get('vendedores', {})
        
        # Análise geral da equipe
        nota_media_equipe = team_stats.get('nota_media_equipe', 0)
        taxa_conversao_media = team_stats.get('taxa_conversao_media', 0)
        
        if nota_media_equipe < 6:
            doc.add_paragraph("🚨 URGENTE: A nota média da equipe está abaixo de 6. Recomenda-se treinamento intensivo em técnicas básicas de vendas.")
        elif nota_media_equipe < 7:
            doc.add_paragraph("⚠️ A equipe tem potencial de melhoria. Foco em treinamentos específicos por vendedor.")
        else:
            doc.add_paragraph("✅ A equipe está performando bem. Manter padrão atual e focar em casos específicos.")
        
        if taxa_conversao_media < 20:
            doc.add_paragraph("📈 Taxa de conversão baixa. Implementar treinamento em técnicas de fechamento e identificação de oportunidades.")
        
        # Recomendações por vendedor
        doc.add_heading('👤 Por Vendedor', level=2)
        
        for vendor_name, vendor_info in vendors.items():
            insights = vendor_info.get('insights', {})
            
            vendor_rec_para = doc.add_paragraph()
            vendor_rec_para.add_run(f"{vendor_name}: ").bold = True
            
            nivel_performance = insights.get('nivel_performance', '')
            if nivel_performance == 'Precisa Melhorar':
                vendor_rec_para.add_run("Prioridade alta para treinamento. ")
            elif nivel_performance == 'Regular':
                vendor_rec_para.add_run("Acompanhamento próximo e treinamento específico. ")
            elif nivel_performance == 'Bom':
                vendor_rec_para.add_run("Manter performance e trabalhar pontos específicos. ")
            else:
                vendor_rec_para.add_run("Excelente performance, pode ser mentor para outros. ")
            
            recomendacao = insights.get('recomendacao_principal', '')
            if recomendacao:
                vendor_rec_para.add_run(recomendacao)
        
        # Próximos passos para gestão
        doc.add_heading('🎯 Próximos Passos para Gestão', level=2)
        
        next_steps = [
            "Implementar plano de treinamento baseado nas prioridades identificadas",
            "Agendar reuniões individuais com vendedores que precisam de atenção",
            "Criar programa de mentoria com os melhores vendedores",
            "Monitorar progresso mensalmente com nova análise de ligações",
            "Desenvolver scripts para objeções mais comuns identificadas"
        ]
        
        for step in next_steps:
            doc.add_paragraph(f"• {step}", style='List Bullet')
        
        # Rodapé
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_para.add_run("Relatório gerado automaticamente pelo VVN AI Analyzer").italic = True
        footer_para.add_run(f"\\n{datetime.now().strftime('%d/%m/%Y %H:%M')}").italic = True

